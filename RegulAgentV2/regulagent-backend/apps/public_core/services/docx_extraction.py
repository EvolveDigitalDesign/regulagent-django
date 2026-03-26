from __future__ import annotations

import base64
import io
import json
import os
import time
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import logging

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "bmp", "tiff"}

CONTENT_TYPE_MAP = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "bmp": "image/bmp",
    "tiff": "image/tiff",
}

MODEL_VISION = os.getenv("OPENAI_VISION_MODEL", "gpt-4o")


@dataclass
class DocxExtractionResult:
    json_data: dict
    model_tag: str
    errors: list
    image_analyses: list = field(default_factory=list)


def extract_text_from_docx(file_path: str) -> Tuple[str, List[dict]]:
    """
    Parse paragraphs and tables from a .docx file.

    Returns (concatenated_text, table_data_list) where concatenated_text is
    truncated to 30,000 chars and table_data_list contains one dict per table.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("extract_text_from_docx: python-docx not installed")
        return "", []

    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error("extract_text_from_docx: failed to open %s: %s", file_path, e)
        return "", []

    text_parts: List[str] = []
    table_data: List[dict] = []

    try:
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                text_parts.append(t)
    except Exception as e:
        logger.warning("extract_text_from_docx: paragraph extraction error: %s", e)

    try:
        for table_index, table in enumerate(doc.tables):
            rows = table.rows
            if not rows:
                continue
            headers = [cell.text.strip() for cell in rows[0].cells]
            records: List[dict] = []
            for row in rows[1:]:
                values = [cell.text.strip() for cell in row.cells]
                if headers and any(v for v in values):
                    record = dict(zip(headers, values))
                    records.append(record)
            if records:
                table_data.append({"table_index": table_index, "headers": headers, "rows": records})
                # Include table text in concatenated output
                text_parts.append(f"[Table {table_index}] " + " | ".join(headers))
                for record in records:
                    text_parts.append(" | ".join(str(v) for v in record.values()))
    except Exception as e:
        logger.warning("extract_text_from_docx: table extraction error: %s", e)

    text = "\n".join(text_parts)
    if len(text) > 30000:
        text = text[:30000]

    return text, table_data


def analyze_docx_images(file_path: str, api_number: str) -> List[dict]:
    """
    Extract and analyze images embedded in a .docx file.

    Opens the docx as a ZIP archive, reads files from word/media/,
    and sends each image to OpenAI Vision API for classification and extraction.
    Returns a list of analysis dicts; no raw image bytes are stored.
    """
    from .openai_extraction import _openai_client

    results: List[dict] = []

    try:
        zf = zipfile.ZipFile(file_path, "r")
    except Exception as e:
        logger.error("analyze_docx_images: failed to open %s as ZIP: %s", file_path, e)
        return results

    media_files = [
        name for name in zf.namelist()
        if name.startswith("word/media/") and name.split(".")[-1].lower() in IMAGE_EXTENSIONS
    ]
    logger.info("analyze_docx_images: found %d image(s) in %s", len(media_files), file_path)

    try:
        client = _openai_client()
    except Exception as e:
        logger.error("analyze_docx_images: cannot create OpenAI client: %s", e)
        zf.close()
        return results

    for image_index, media_name in enumerate(media_files):
        ext = media_name.split(".")[-1].lower()
        content_type = CONTENT_TYPE_MAP.get(ext, "image/jpeg")
        try:
            image_bytes = zf.read(media_name)
            size_bytes = len(image_bytes)
            b64 = base64.b64encode(image_bytes).decode("utf-8")
            data_url = f"data:{content_type};base64,{b64}"

            # Step 1: Classify the image
            classification = _vision_classify(client, data_url, image_index)

            # Step 2: Structured extraction or rich summary depending on class
            structured_data: Dict[str, Any] = {}
            text_summary = ""

            if classification in ("schematic", "cbl_log"):
                structured_data = _vision_extract_structured(client, data_url, classification, image_index)
            else:
                text_summary = _vision_summarize(client, data_url, image_index)

            results.append({
                "image_index": image_index,
                "classification": classification,
                "structured_data": structured_data,
                "text_summary": text_summary,
                "content_type": content_type,
                "size_bytes": size_bytes,
            })
        except Exception as e:
            logger.warning(
                "analyze_docx_images: error processing image %d (%s): %s",
                image_index,
                media_name,
                e,
            )
            results.append({
                "image_index": image_index,
                "classification": "error",
                "structured_data": {},
                "text_summary": "",
                "content_type": content_type,
                "size_bytes": 0,
                "error": str(e),
            })

    zf.close()
    return results


def _vision_classify(client, data_url: str, image_index: int) -> str:
    prompt = (
        "Classify this image from an oil/gas P&A packet. "
        "Respond with exactly one word from: schematic, cbl_log, map, permit, photo, unknown"
    )
    try:
        resp = client.chat.completions.create(
            model=MODEL_VISION,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                }
            ],
            max_tokens=20,
            temperature=0,
        )
        label = (resp.choices[0].message.content or "").strip().lower().split()[0]
        valid = {"schematic", "cbl_log", "map", "permit", "photo", "unknown"}
        return label if label in valid else "unknown"
    except Exception as e:
        logger.warning("_vision_classify: image_index=%d error=%s", image_index, e)
        return "unknown"


def _vision_extract_structured(client, data_url: str, classification: str, image_index: int) -> Dict[str, Any]:
    prompt = (
        "Extract from this well schematic/CBL: casing depths, cement bond quality zones, "
        "formation intervals, any annotations with measurements. "
        "Return valid JSON only."
    )
    try:
        resp = client.chat.completions.create(
            model=MODEL_VISION,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                }
            ],
            max_tokens=1000,
            temperature=0,
        )
        content = (resp.choices[0].message.content or "").strip()
        # Strip markdown fences if present
        if content.startswith("```"):
            content = content.split("```", 2)[1]
            if content.startswith("json"):
                content = content[4:]
            content = content.rstrip("`").strip()
        return json.loads(content)
    except Exception as e:
        logger.warning("_vision_extract_structured: image_index=%d error=%s", image_index, e)
        return {}


def _vision_summarize(client, data_url: str, image_index: int) -> str:
    prompt = (
        "Describe this image from an oil/gas P&A packet. "
        "Include any visible depths, readings, annotations, labels, or measurements."
    )
    try:
        resp = client.chat.completions.create(
            model=MODEL_VISION,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                }
            ],
            max_tokens=500,
            temperature=0,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        logger.warning("_vision_summarize: image_index=%d error=%s", image_index, e)
        return ""


def extract_pa_procedure_from_docx(file_path: str, api_number: str) -> DocxExtractionResult:
    """
    Extract P&A procedure data from a .docx file.

    Combines text extraction, OpenAI Responses API JSON extraction, and
    image analysis. Returns a DocxExtractionResult with merged image analyses.
    """
    from .openai_extraction import _openai_client, _load_prompt, MODEL_PRIMARY

    errors: List[str] = []

    # 1. Extract text and table data
    text, table_data = extract_text_from_docx(file_path)
    if not text:
        logger.warning("extract_pa_procedure_from_docx: no text extracted from %s", file_path)

    # 2. Load the pa_procedure prompt
    try:
        base_prompt = _load_prompt("pa_procedure")
    except Exception as e:
        logger.error("extract_pa_procedure_from_docx: failed to load prompt: %s", e)
        errors.append(f"prompt_load_error: {e}")
        base_prompt = "Extract P&A procedure data. Return JSON."

    # 3. Build system prompt with table context
    table_context = ""
    if table_data:
        try:
            table_context = "\n\nTABLE DATA:\n" + json.dumps(table_data, ensure_ascii=False)[:5000]
        except Exception:
            pass

    system_prompt = base_prompt + " Return only valid JSON." + table_context

    # 4. Call OpenAI Responses API
    json_data: dict = {}
    model_tag = MODEL_PRIMARY
    last_err = None

    try:
        client = _openai_client()
    except Exception as e:
        logger.error("extract_pa_procedure_from_docx: cannot create OpenAI client: %s", e)
        errors.append(f"openai_client_error: {e}")
        client = None

    if client and text:
        for attempt in range(3):
            logger.info(
                "extract_pa_procedure_from_docx: attempt=%d file=%s model=%s",
                attempt + 1,
                file_path,
                model_tag,
            )
            try:
                inputs = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": system_prompt},
                            {"type": "input_text", "text": text},
                        ],
                    }
                ]
                resp = client.responses.create(
                    model=model_tag,
                    input=inputs,
                    text={"format": {"type": "json_object"}},
                    max_output_tokens=8000,
                    temperature=0,
                )

                # Extract text content from Responses API output
                content = ""
                if hasattr(resp, "output") and resp.output:
                    try:
                        content = "".join(
                            (
                                block.get("text", "")
                                if isinstance(block, dict)
                                else (getattr(block, "text", "") or "")
                            )
                            for item in resp.output
                            for block in (
                                item.get("content", []) if isinstance(item, dict)
                                else (getattr(item, "content", []) or [])
                            )
                            if (
                                block.get("type") if isinstance(block, dict)
                                else getattr(block, "type", None)
                            ) == "output_text"
                        )
                    except Exception as parse_e:
                        logger.warning(
                            "extract_pa_procedure_from_docx: failed to parse output blocks: %s",
                            parse_e,
                        )
                if not content and hasattr(resp, "output_text"):
                    content = resp.output_text or ""

                if not content or content.strip() in ("{}", "[]", "null", "None", ""):
                    raise ValueError("EMPTY_JSON_RESPONSE")

                logger.info("extract_pa_procedure_from_docx: received json length=%d", len(content))
                json_data = json.loads(content)
                break
            except Exception as e:
                last_err = str(e)
                logger.warning(
                    "extract_pa_procedure_from_docx: error attempt=%d err=%s",
                    attempt + 1,
                    last_err,
                )
                time.sleep(0.5 if attempt == 0 else 2.0)

        if not json_data and last_err:
            errors.append(f"extraction_error: {last_err}")
            logger.error("extract_pa_procedure_from_docx: failed after retries err=%s", last_err)

    # 5. Analyze images
    image_analyses: List[dict] = []
    try:
        image_analyses = analyze_docx_images(file_path, api_number)
        logger.info(
            "extract_pa_procedure_from_docx: image_analyses count=%d", len(image_analyses)
        )
    except Exception as e:
        logger.warning("extract_pa_procedure_from_docx: image analysis error: %s", e)
        errors.append(f"image_analysis_error: {e}")

    # 6. Merge image analyses into json_data
    if image_analyses:
        json_data["embedded_images"] = image_analyses

    return DocxExtractionResult(
        json_data=json_data,
        model_tag=model_tag,
        errors=errors,
        image_analyses=image_analyses,
    )
